"""
文档组装与质检 — 合并全文、生成目录、质量检查、导出多格式。
"""

import logging
import re
from pathlib import Path
from datetime import datetime

from ..llm_provider import LLMClient, Settings
from ..planner.outline import Outline
from ..generator.writer import MemoryPack
from ..renderer.pdf_book import PDFBookRenderer

logger = logging.getLogger(__name__)


class DocAssembler:
    """将各节内容组装为完整文档，并进行质检和导出。"""

    def __init__(self, llm: LLMClient, settings: Settings):
        self.llm = llm
        self.settings = settings
        self.output_cfg = settings.output
        self.render_cfg = settings.render or {}
        project_root = Path(__file__).resolve().parents[2]
        output_dir_cfg = Path(settings.output.get("output_dir", "data/output"))
        self.output_dir = output_dir_cfg if output_dir_cfg.is_absolute() else (project_root / output_dir_cfg)

    def assemble(
        self,
        outline: Outline,
        sections: dict[str, str],
        memory: MemoryPack = None,
    ) -> str:
        parts = []

        # 标题
        parts.append(f"# {outline.topic}\n")
        parts.append(f"> 生成日期：{datetime.now().strftime('%Y-%m-%d')}\n")

        # 目录
        if self.output_cfg.get("include_toc", True):
            parts.append(self._build_toc(outline))

        # 正文
        current_chapter = None
        for section in outline.sections:
            if section.chapter != current_chapter:
                current_chapter = section.chapter
                parts.append(f"\n## {current_chapter}\n")
            content = sections.get(section.id, "")
            if content:
                parts.append(f"### {section.title}\n")
                parts.append(content)
                parts.append("")

        # 术语表
        if self.output_cfg.get("include_glossary", True) and memory and memory.glossary:
            parts.append(self._build_glossary(memory))

        full_text = "\n".join(parts)
        char_count = len(re.sub(r'\s', '', full_text))
        logger.info(f"文档组装完成: {char_count} 字（含标点空格: {len(full_text)}）")
        return full_text

    def quality_check(self, text: str, outline: Outline) -> dict:
        """基础质量检查。"""
        issues = []
        char_count = len(re.sub(r'\s', '', text))
        min_chars = self.settings.pipeline.get("min_chars", 45000)

        if char_count < min_chars:
            issues.append(f"字数不足: {char_count} < {min_chars}")

        # 重复段落检测
        paragraphs = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 50]
        seen = set()
        for p in paragraphs:
            key = p[:100]
            if key in seen:
                issues.append(f"发现重复段落: {key[:50]}...")
            seen.add(key)

        # 章节完整性
        for ch in outline.chapters:
            if ch["title"] not in text:
                issues.append(f"缺失章节: {ch['title']}")

        return {
            "char_count": char_count,
            "paragraph_count": len(paragraphs),
            "chapter_count": text.count("\n## "),
            "section_count": text.count("\n### "),
            "issues": issues,
            "passed": len(issues) == 0,
        }

    def llm_review(self, text: str) -> str:
        """LLM 全文审核。"""
        sample = text[:5000] + "\n...\n" + text[-3000:]
        system = """你是一位资深编辑。请审核以下文档的前后部分，检查：
1. 逻辑连贯性 - 前后内容是否有矛盾
2. 信息准确性 - 是否有明显的技术错误
3. 可读性 - 是否有表述不清或冗余的地方
4. AI痕迹 - 是否仍有明显的AI生成特征

输出格式：问题列表（如果没问题就说"通过审核"）"""

        return self.llm.call("review", [
            {"role": "user", "content": f"请审核：\n\n{sample}"}
        ], system=system)

    def export(self, text: str, topic: str, render_options: dict | None = None) -> list[str]:
        """导出为指定格式。"""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', topic)[:50]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"{safe_name}_{timestamp}"
        exported = []

        formats = self.output_cfg.get("formats", ["markdown"])

        md_path = None
        if "markdown" in formats:
            md_path = self.output_dir / f"{base_name}.md"
            md_path.write_text(text, encoding="utf-8")
            exported.append(str(md_path))
            logger.info(f"导出 Markdown: {md_path}")

        if "docx" in formats:
            docx_path = self._export_docx(text, base_name)
            if docx_path:
                exported.append(docx_path)

        if "pdf" in formats:
            pdf_path = self._export_pdf(text, base_name)
            if pdf_path:
                exported.append(pdf_path)

        if render_options and render_options.get("to_pdf"):
            if md_path is None:
                md_path = self.output_dir / f"{base_name}.md"
                md_path.write_text(text, encoding="utf-8")
                exported.append(str(md_path))
                logger.info(f"导出 Markdown: {md_path}")
            custom_pdf = self.render_pdf(
                markdown_path=md_path,
                title=render_options.get("title") or topic,
                subtitle=render_options.get("subtitle", ""),
                author=render_options.get("author", ""),
                basename=render_options.get("basename") or base_name,
                pdf_theme=render_options.get("pdf_theme"),
                illustration=render_options.get("illustration"),
            )
            exported.append(custom_pdf)

        return exported

    def render_pdf(
        self,
        markdown_path: str | Path,
        title: str,
        subtitle: str = "",
        author: str = "",
        basename: str | None = None,
        pdf_theme: dict | None = None,
        illustration: dict | None = None,
    ) -> str:
        markdown_path = Path(markdown_path)
        if basename:
            pdf_name = f"{basename}.pdf"
        else:
            pdf_name = markdown_path.with_suffix(".pdf").name
        output_path = self.output_dir / pdf_name

        render_cfg = dict(self.render_cfg)
        if pdf_theme:
            render_pdf_cfg = dict(render_cfg.get("pdf", {}))
            current_theme = dict(render_pdf_cfg.get("theme", {}))
            current_theme.update(pdf_theme)
            render_pdf_cfg["theme"] = current_theme
            render_cfg["pdf"] = render_pdf_cfg

        renderer = PDFBookRenderer(render_cfg)
        markdown_text = markdown_path.read_text(encoding="utf-8", errors="ignore")
        result = renderer.render(
            markdown_text,
            output_path,
            title=title,
            subtitle=subtitle,
            author=author,
            illustration=illustration,
        )
        logger.info(f"导出出版级 PDF: {result}")
        return result

    def _build_toc(self, outline: Outline) -> str:
        lines = ["\n## 目录\n"]
        for ch in outline.chapters:
            lines.append(f"- **{ch['title']}**")
            for sec in outline.sections:
                if sec.chapter == ch["title"]:
                    lines.append(f"  - {sec.title}")
        lines.append("")
        return "\n".join(lines)

    def _build_glossary(self, memory: MemoryPack) -> str:
        lines = ["\n## 术语表\n"]
        for term, defn in sorted(memory.glossary.items()):
            lines.append(f"- **{term}**：{defn}")
        return "\n".join(lines)

    def _export_docx(self, text: str, base_name: str) -> str | None:
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            for line in text.split("\n"):
                line = line.strip()
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=0)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line:
                    doc.add_paragraph(line)
            path = self.output_dir / f"{base_name}.docx"
            doc.save(str(path))
            logger.info(f"导出 DOCX: {path}")
            return str(path)
        except ImportError:
            logger.warning("python-docx 未安装，跳过 DOCX 导出")
            return None

    def _export_pdf(self, text: str, base_name: str) -> str | None:
        try:
            import markdown
            from weasyprint import HTML
            html = markdown.markdown(text, extensions=["tables", "fenced_code"])
            styled = f"""<html><head><style>
body {{ font-family: "Microsoft YaHei", sans-serif; max-width: 800px; margin: 0 auto; padding: 40px; line-height: 1.8; }}
h1 {{ font-size: 28px; }} h2 {{ font-size: 22px; margin-top: 30px; }} h3 {{ font-size: 18px; }}
code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; }}
pre {{ background: #f5f5f5; padding: 16px; border-radius: 6px; overflow-x: auto; }}
</style></head><body>{html}</body></html>"""
            path = self.output_dir / f"{base_name}.pdf"
            HTML(string=styled).write_pdf(str(path))
            logger.info(f"导出 PDF: {path}")
            return str(path)
        except ImportError:
            logger.warning("weasyprint 或 markdown 未安装，跳过 PDF 导出")
            return None
